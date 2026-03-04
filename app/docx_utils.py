from __future__ import annotations

from pathlib import Path


def markdown_to_docx(markdown_text: str, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception as exc:
        raise RuntimeError("python-docx not installed") from exc

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:])
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:])
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(13)
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:])
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        else:
            doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
