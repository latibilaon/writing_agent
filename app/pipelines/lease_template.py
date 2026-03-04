from __future__ import annotations

import datetime
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from ..converter import convert_tree, load_markdown_bundle
from ..openrouter_client import OpenRouterClient


@dataclass
class LeaseTemplateRequest:
    tenant_name: str
    property_address: str
    jurisdiction: str
    issues: str
    demands: str
    template_path: Path
    model: str
    materials_dir: Path
    converted_dir: Path
    output_dir: Path
    skip_convert: bool = False


def _slug(value: str) -> str:
    s = re.sub(r"\s+", "_", value.strip())
    return re.sub(r"[^A-Za-z0-9_\-]", "", s) or "Tenant"


def _extract_paragraphs(doc: Document) -> list[dict]:
    items = []
    idx = 0
    for para in doc.paragraphs:
        items.append({"idx": idx, "ref": para, "text": para.text})
        idx += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    items.append({"idx": idx, "ref": para, "text": para.text})
                    idx += 1
    return items


def _apply_text_with_bold(para, new_text: str):
    for run in list(para.runs):
        para._p.remove(run._r)
    parts = re.split(r"\*\*(.*?)\*\*", new_text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (i % 2 == 1)


def _analysis_prompt(req: LeaseTemplateRequest, bundle: str) -> str:
    return f"""Extract concise landlord breach points and related legal framing from materials.
Output English markdown bullet list.

Tenant issues:
{req.issues}

Materials:
{bundle if bundle else '[No materials]'}
"""


def _rewrite_prompt(req: LeaseTemplateRequest, items: list[dict], analysis: str) -> str:
    numbered = "\n".join(f"[P{it['idx']:03d}] {it['text']}" for it in items if it["text"].strip())
    return f"""You are editing a lease-termination template.
Return EVERY paragraph in exact [P###] format. Do not add or remove paragraphs.
Only rewrite where needed for this case.
Wrap modified/newly emphasized text with **double asterisks**.
No commentary outside [P###] lines.

Case:
- Tenant: {req.tenant_name}
- Address: {req.property_address}
- Jurisdiction: {req.jurisdiction}
- Demands: {req.demands}
- Issues: {req.issues}

Contract analysis:
{analysis}

Template paragraphs:
{numbered}
"""


def _parse_rewrite(response: str, original_items: list[dict]) -> dict[int, str]:
    pattern = re.compile(r"\[P(\d+)\]\s*(.*?)(?=\[P\d+\]|$)", re.DOTALL)
    parsed = {}
    for m in pattern.finditer(response):
        idx = int(m.group(1))
        txt = m.group(2).strip()
        parsed[idx] = txt

    originals = {x["idx"]: x["text"] for x in original_items}
    changed = {}
    for idx, txt in parsed.items():
        if idx in originals and txt != originals[idx]:
            changed[idx] = txt
    return changed


def run_lease_template_pipeline(req: LeaseTemplateRequest, client: OpenRouterClient, logger):
    logger.head("=" * 60)
    logger.head("STEP 1  Contract Materials -> Markdown")
    logger.head("=" * 60)
    req.converted_dir.mkdir(parents=True, exist_ok=True)
    req.output_dir.mkdir(parents=True, exist_ok=True)
    if not req.skip_convert:
        convert_tree(req.materials_dir, req.converted_dir, logger)
    else:
        logger.warn("[SKIP] conversion skipped by user")

    bundle = load_markdown_bundle(req.converted_dir)

    logger.head("=" * 60)
    logger.head("STEP 2  Contract Analysis")
    logger.head("=" * 60)
    analysis, usage = client.chat(req.model, _analysis_prompt(req, bundle), x_title="Uoffer Portable Template Analysis")
    logger.info(f"[INFO] analysis usage: prompt={usage.get('prompt_tokens')} completion={usage.get('completion_tokens')}")

    logger.head("=" * 60)
    logger.head("STEP 3  Template Rewrite")
    logger.head("=" * 60)
    doc = Document(str(req.template_path))
    items = _extract_paragraphs(doc)
    rewritten, usage2 = client.chat(req.model, _rewrite_prompt(req, items, analysis), x_title="Uoffer Portable Template Rewrite")
    logger.info(f"[INFO] rewrite usage: prompt={usage2.get('prompt_tokens')} completion={usage2.get('completion_tokens')}")

    changes = _parse_rewrite(rewritten, items)
    applied = 0
    for item in items:
        if item["idx"] in changes:
            _apply_text_with_bold(item["ref"], changes[item["idx"]])
            applied += 1

    logger.head("=" * 60)
    logger.head("STEP 4  Save Output")
    logger.head("=" * 60)
    day = datetime.date.today().strftime("%Y%m%d")
    slug = _slug(req.tenant_name)
    out = req.output_dir / f"Lease_Template_{slug}_{day}.docx"
    doc.save(str(out))
    logger.ok(f"[OK] Saved: {out} ({applied} paragraphs updated)")
    return {"docx": out, "updated_paragraphs": applied}
